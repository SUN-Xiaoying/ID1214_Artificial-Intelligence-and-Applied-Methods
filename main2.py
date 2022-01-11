import os
import time
import paddlehub as hub
from helpers import helper
from docx import Document

'''
Translte folder
'''
def doc_translate(input_file, output_file, modelname):
    # Extract Paragraph
    text = helper.get_paragraphs_text(input_file)
    # Improt module
    model = hub.Module(name=modelname, beam_size=5)
    # Create new Document object 
    new_document = Document()
    for item in text:
        src_style = item[0]
        src_texts = [item[1]]
        print('标题层级：', src_style)
        print(src_texts)
        n_best = 1  # One input matches one output
        trg_texts = model.predict(src_texts, n_best=n_best)
        print(trg_texts)
        # Match title
        if src_style == 'Title':
            new_document.add_heading(trg_texts, 0)
        elif src_style[:7:] == 'Heading':
            new_document.add_heading(trg_texts, level=int(src_style[-1]))
        else:
            new_document.add_paragraph(trg_texts)
    new_document.save(output_file)

word_folder = './docx'

files = os.listdir(word_folder)  
for fi in files:
    input_file = os.path.join(word_folder,fi)
    output_file = os.path.join(word_folder,'translate_' + os.path.basename(input_file))
    doc_translate(input_file, output_file, 'transformer_zh-en')
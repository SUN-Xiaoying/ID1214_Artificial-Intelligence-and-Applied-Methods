import os
import time
import paddlehub as hub
from helpers import helper
from docx import Document

'''
Translte single docx file
'''

text = helper.get_paragraphs_text('docx/zh-demo.docx')

model = hub.Module(name='transformer_zh-en', beam_size=5)
new_document = Document()
# starting time
start = time.time()

for item in text:
    src_style = item[0]
    src_texts = [item[1]]
    print('Heading level: ', src_style)
    print(src_texts)
    n_best = 1  # 每个输入样本的输出候选句子数量
    trg_texts = model.predict(src_texts, n_best=n_best)
    print(trg_texts)
    if src_style == 'Title':
        new_document.add_heading(trg_texts, 0)
    elif src_style[:7:] == 'Heading':
        new_document.add_heading(trg_texts, level=int(src_style[-1]))
    else:
        new_document.add_paragraph(trg_texts)

# end time
end = time.time()

print('Time: ', end - start)  

new_document.save('docx/translate_demo.docx')

# Clean untranslated result
document = Document('docx/translate_demo.docx')
# Display the content of each paragraph
helper.replace_text(document, 'No', '')
# Save cleaned documents
document.save('docx/translate_clean_demo.docx')
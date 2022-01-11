from docx import Document

def get_paragraphs_text(path):
    """
    param path: word.docx
    return list
    """
    document = Document(path)
    all_paragraphs = document.paragraphs
    paragraphs_text = []
    for paragraph in all_paragraphs:
        # 拼接一个list,包括段落的结构和内容
        paragraphs_text.append([paragraph.style.name,paragraph.text])
    return paragraphs_text


'''
global content replacement
'''
def replace_text(doc, old_text, new_text):
    # iterate each paragraph
    for p in doc.paragraphs:
        # If search content in the paragraph
        if old_text in p.text:
            # Use 'runs' to replace the content without changing the style

            for run in p.runs:
                if old_text in run.text:
                    run.text = run.text.replace(old_text, new_text)

from docx import Document
from docx.shared import Inches

document = Document()

document.add_heading('标题', 0)

p = document.add_paragraph('一个简单的段落，有一些')
p.add_run('粗体').bold = True
p.add_run('和一些')
p.add_run('斜体。').italic = True

document.add_heading('一级标题', level=1)
document.add_paragraph('一级标题下的普通文本。')

document.add_heading('二级标题', level=2)
document.add_paragraph('二级标题下的普通文本。')

document.add_heading('三级标题', level=3)
document.add_paragraph('三级标题下的普通文本。')

document.add_paragraph('强引用', style='Intense Quote')

document.add_paragraph(
    '无序列表第一项', style='List Bullet'
)
document.add_paragraph(
    '有序列表第一项', style='List Number'
)

document.add_picture('./imgs/sign.png', width=Inches(1.25))

records = (
    (3, '101', 'Spam'),
    (7, '422', 'Eggs'),
    (4, '631', 'Spam, spam, eggs, and spam')
)

table = document.add_table(rows=1, cols=3)
hdr_cells = table.rows[0].cells
hdr_cells[0].text = '表格第一列'
hdr_cells[1].text = '表格第二列'
hdr_cells[2].text = '表格第三列'
for qty, id, desc in records:
    row_cells = table.add_row().cells
    row_cells[0].text = str(qty)
    row_cells[1].text = id
    row_cells[2].text = desc

document.add_page_break()

document.save('zh-demo.docx')